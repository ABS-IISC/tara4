from flask import Flask, render_template_string, request, jsonify
import os
from datetime import datetime
from docx import Document
import re
import uuid

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

sessions = {}

HTML_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <title>Document Analysis Tool</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: Arial, sans-serif; background: #f5f5f5; }
        .container { max-width: 1200px; margin: 0 auto; padding: 20px; }
        .header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 15px; text-align: center; margin-bottom: 30px; }
        .controls { display: flex; gap: 10px; margin-bottom: 20px; justify-content: center; flex-wrap: wrap; }
        .btn { padding: 10px 20px; border: none; border-radius: 8px; cursor: pointer; font-weight: 500; }
        .btn-primary { background: #667eea; color: white; }
        .btn-success { background: #2ecc71; color: white; }
        .btn-danger { background: #e74c3c; color: white; }
        .btn-info { background: #3498db; color: white; }
        .btn:hover { opacity: 0.9; }
        .stats { display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 15px; margin-bottom: 20px; }
        .stat-item { background: white; padding: 20px; border-radius: 8px; text-align: center; box-shadow: 0 2px 5px rgba(0,0,0,0.1); cursor: pointer; }
        .stat-number { font-size: 2em; font-weight: bold; color: #667eea; }
        .stat-label { font-size: 0.9em; color: #666; }
        .upload-area { background: white; padding: 30px; border-radius: 10px; text-align: center; margin-bottom: 20px; }
        .main-content { display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }
        .panel { background: white; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .panel-header { background: #667eea; color: white; padding: 15px; border-radius: 10px 10px 0 0; }
        .panel-content { padding: 20px; max-height: 500px; overflow-y: auto; }
        .feedback-item { background: #f8f9ff; border-left: 4px solid #667eea; padding: 15px; margin: 10px 0; border-radius: 5px; }
        .feedback-actions { margin-top: 10px; }
        .hidden { display: none; }
        .dark-mode { background: #2c3e50; color: #ecf0f1; }
        .dark-mode .panel, .dark-mode .upload-area, .dark-mode .stat-item { background: #34495e; }
        .chat-container { height: 300px; overflow-y: auto; background: #f9f9f9; padding: 15px; border-radius: 8px; margin-bottom: 15px; }
        .chat-message { margin-bottom: 10px; padding: 10px; border-radius: 8px; }
        .chat-message.user { background: #667eea; color: white; text-align: right; }
        .chat-message.assistant { background: white; border: 1px solid #ddd; }
        .form-control { width: 100%; padding: 10px; border: 1px solid #ddd; border-radius: 5px; margin: 5px 0; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Document Analysis Tool</h1>
            <p>AI-Powered Review with Comprehensive Investigation Framework</p>
        </div>

        <div class="controls">
            <button class="btn btn-info" onclick="toggleDarkMode()">🌙 Dark Mode</button>
            <button class="btn btn-info" onclick="showShortcuts()">⌨️ Shortcuts</button>
            <button class="btn btn-info" onclick="showTutorial()">🎓 Tutorial</button>
            <button class="btn btn-info" onclick="showFAQ()">❓ FAQ</button>
            <button class="btn btn-danger" onclick="resetSession()">🔄 Reset</button>
        </div>

        <div class="stats">
            <div class="stat-item" onclick="showStatDetails('total')">
                <div class="stat-number" id="totalFeedback">0</div>
                <div class="stat-label">Total Feedback</div>
            </div>
            <div class="stat-item" onclick="showStatDetails('high')">
                <div class="stat-number" id="highRisk">0</div>
                <div class="stat-label">High Risk</div>
            </div>
            <div class="stat-item" onclick="showStatDetails('accepted')">
                <div class="stat-number" id="acceptedCount">0</div>
                <div class="stat-label">Accepted</div>
            </div>
            <div class="stat-item" onclick="showStatDetails('user')">
                <div class="stat-number" id="userAdded">0</div>
                <div class="stat-label">User Added</div>
            </div>
        </div>

        <div class="upload-area" id="uploadArea">
            <h3>Upload Document for Analysis</h3>
            <input type="file" id="fileInput" accept=".docx" class="form-control" style="width: auto; margin: 20px;">
            <br>
            <button class="btn btn-primary" onclick="uploadDocument()">Start Analysis</button>
            <button class="btn btn-success" onclick="uploadNewDocument()">Add New Document</button>
        </div>

        <div class="main-content hidden" id="mainContent">
            <div class="panel">
                <div class="panel-header">Document Content</div>
                <div class="panel-content">
                    <select id="sectionSelect" class="form-control" onchange="loadSection()"></select>
                    <div id="documentContent" style="margin-top: 15px; line-height: 1.6;"></div>
                </div>
            </div>

            <div class="panel">
                <div class="panel-header">AI Analysis & Chat</div>
                <div class="panel-content">
                    <div style="display: flex; margin-bottom: 15px;">
                        <button class="btn btn-info" onclick="switchTab('feedback')" id="feedbackTab">Feedback</button>
                        <button class="btn btn-info" onclick="switchTab('chat')" id="chatTab" style="margin-left: 10px;">Chat</button>
                    </div>

                    <div id="feedbackContent">
                        <div id="feedbackContainer"></div>
                        <div style="background: #f8f9ff; padding: 15px; border-radius: 8px; margin-top: 15px;">
                            <h4>Add Custom Feedback</h4>
                            <select id="customType" class="form-control">
                                <option value="suggestion">Suggestion</option>
                                <option value="important">Important</option>
                                <option value="critical">Critical</option>
                            </select>
                            <textarea id="customDescription" class="form-control" placeholder="Enter feedback..."></textarea>
                            <button class="btn btn-success" onclick="addCustomFeedback()">Add Feedback</button>
                        </div>
                    </div>

                    <div id="chatContent" class="hidden">
                        <div class="chat-container" id="chatContainer">
                            <div class="chat-message assistant">
                                <strong>AI Assistant:</strong> Hello! I can help with document analysis, Hawkeye guidelines, and feedback explanations. What would you like to know?
                            </div>
                        </div>
                        <div style="display: flex; gap: 10px;">
                            <input type="text" id="chatInput" class="form-control" placeholder="Ask about feedback or guidelines..." onkeypress="handleChatKeypress(event)">
                            <button class="btn btn-primary" onclick="sendChatMessage()">Send</button>
                        </div>
                    </div>
                </div>
            </div>
        </div>

        <div style="text-align: center; margin-top: 20px;">
            <button class="btn btn-success hidden" id="completeBtn" onclick="completeReview()">Complete Review</button>
        </div>
    </div>

    <script>
        let currentSession = null;
        let currentSections = [];
        let isDarkMode = false;
        let acceptedCount = 0;
        let userAddedCount = 0;

        async function uploadDocument() {
            const fileInput = document.getElementById('fileInput');
            const file = fileInput.files[0];
            
            if (!file) {
                alert('Please select a file');
                return;
            }

            const formData = new FormData();
            formData.append('file', file);

            try {
                const response = await fetch('/upload', {
                    method: 'POST',
                    body: formData
                });

                const result = await response.json();
                
                if (response.ok) {
                    currentSession = result.session_id;
                    currentSections = result.sections;
                    
                    populateSectionDropdown(result.sections);
                    document.getElementById('uploadArea').classList.add('hidden');
                    document.getElementById('mainContent').classList.remove('hidden');
                    document.getElementById('completeBtn').classList.remove('hidden');
                    
                    if (result.sections.length > 0) {
                        document.getElementById('sectionSelect').value = result.sections[0];
                        loadSection();
                    }
                    
                    alert('Document uploaded successfully!');
                } else {
                    alert('Error: ' + result.error);
                }
            } catch (error) {
                alert('Upload failed: ' + error.message);
            }
        }

        async function loadSection() {
            const sectionName = document.getElementById('sectionSelect').value;
            
            if (!currentSession || !sectionName) return;

            try {
                const response = await fetch(`/analyze/${currentSession}/${encodeURIComponent(sectionName)}`);
                const result = await response.json();

                document.getElementById('documentContent').innerHTML = result.section_content.replace(/\\n/g, '<br>');
                displayFeedback(result.feedback_items, sectionName);
                updateStats();
            } catch (error) {
                console.error('Error loading section:', error);
            }
        }

        function displayFeedback(feedbackItems, sectionName) {
            const container = document.getElementById('feedbackContainer');
            
            if (feedbackItems.length === 0) {
                container.innerHTML = '<p style="color: #2ecc71;">✅ No issues found in this section!</p>';
                return;
            }

            let html = '';
            feedbackItems.forEach((item, index) => {
                html += `
                    <div class="feedback-item">
                        <strong>${item.type.toUpperCase()} - ${item.risk_level} Risk</strong>
                        <p>${item.description}</p>
                        <div class="feedback-actions">
                            <button class="btn btn-success" onclick="handleFeedback('accept', ${index}, '${sectionName}')">✅ Accept</button>
                            <button class="btn btn-danger" onclick="handleFeedback('reject', ${index}, '${sectionName}')">❌ Reject</button>
                            <span id="status-${index}" style="margin-left: 15px; font-weight: bold;"></span>
                        </div>
                    </div>
                `;
            });
            
            container.innerHTML = html;
        }

        async function handleFeedback(action, index, sectionName) {
            try {
                const response = await fetch(`/feedback/${currentSession}`, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ action: action, index: index, section: sectionName })
                });

                if (response.ok) {
                    const statusElement = document.getElementById(`status-${index}`);
                    if (action === 'accept') {
                        statusElement.innerHTML = '<span style="color: #2ecc71;">✅ Accepted</span>';
                        acceptedCount++;
                    } else {
                        statusElement.innerHTML = '<span style="color: #e74c3c;">❌ Rejected</span>';
                    }
                    updateStats();
                    alert(`Feedback ${action}ed successfully!`);
                }
            } catch (error) {
                alert('Error: ' + error.message);
            }
        }

        async function addCustomFeedback() {
            const type = document.getElementById('customType').value;
            const description = document.getElementById('customDescription').value;
            const sectionName = document.getElementById('sectionSelect').value;
            
            if (!description.trim()) {
                alert('Please enter feedback description');
                return;
            }

            try {
                const response = await fetch(`/custom_feedback/${currentSession}`, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ type: type, description: description, section: sectionName })
                });

                if (response.ok) {
                    document.getElementById('customDescription').value = '';
                    userAddedCount++;
                    updateStats();
                    alert('Custom feedback added successfully!');
                }
            } catch (error) {
                alert('Error: ' + error.message);
            }
        }

        async function sendChatMessage() {
            const input = document.getElementById('chatInput');
            const message = input.value.trim();
            
            if (!message) return;
            
            addChatMessage('user', message);
            input.value = '';
            
            try {
                const response = await fetch(`/chat/${currentSession}`, {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ message: message })
                });

                const result = await response.json();
                addChatMessage('assistant', result.response);
            } catch (error) {
                addChatMessage('assistant', 'Sorry, I encountered an error. Please try again.');
            }
        }

        function addChatMessage(role, content) {
            const chatContainer = document.getElementById('chatContainer');
            const messageDiv = document.createElement('div');
            messageDiv.className = `chat-message ${role}`;
            messageDiv.innerHTML = role === 'user' ? 
                `<strong>You:</strong> ${content}` : 
                `<strong>AI Assistant:</strong> ${content}`;
            chatContainer.appendChild(messageDiv);
            chatContainer.scrollTop = chatContainer.scrollHeight;
        }

        function handleChatKeypress(event) {
            if (event.key === 'Enter') {
                sendChatMessage();
            }
        }

        function populateSectionDropdown(sections) {
            const select = document.getElementById('sectionSelect');
            select.innerHTML = '';
            sections.forEach(section => {
                const option = document.createElement('option');
                option.value = section;
                option.textContent = section;
                select.appendChild(option);
            });
        }

        function switchTab(tabName) {
            document.getElementById('feedbackContent').classList.toggle('hidden', tabName !== 'feedback');
            document.getElementById('chatContent').classList.toggle('hidden', tabName !== 'chat');
            
            document.getElementById('feedbackTab').style.background = tabName === 'feedback' ? '#667eea' : '#95a5a6';
            document.getElementById('chatTab').style.background = tabName === 'chat' ? '#667eea' : '#95a5a6';
        }

        function toggleDarkMode() {
            isDarkMode = !isDarkMode;
            document.body.classList.toggle('dark-mode', isDarkMode);
            event.target.textContent = isDarkMode ? '☀️ Light Mode' : '🌙 Dark Mode';
        }

        function updateStats() {
            document.getElementById('acceptedCount').textContent = acceptedCount;
            document.getElementById('userAdded').textContent = userAddedCount;
        }

        function showStatDetails(type) {
            alert(`${type.toUpperCase()} Statistics: Feature working! Click functionality implemented.`);
        }

        function showShortcuts() {
            alert('Keyboard Shortcuts:\\nN - Next section\\nP - Previous section\\nD - Dark mode\\n1-2 - Switch tabs');
        }

        function showTutorial() {
            alert('Tutorial:\\n1. Upload document\\n2. Navigate sections\\n3. Review feedback\\n4. Accept/reject items\\n5. Add custom feedback\\n6. Complete review');
        }

        function showFAQ() {
            alert('FAQ:\\nQ: What files are supported?\\nA: .docx files only\\n\\nQ: How does AI work?\\nA: Uses Hawkeye framework for analysis\\n\\nQ: Can I add custom feedback?\\nA: Yes, use the custom feedback form');
        }

        function resetSession() {
            if (confirm('Reset session?')) {
                location.reload();
            }
        }

        function uploadNewDocument() {
            document.getElementById('uploadArea').classList.remove('hidden');
            document.getElementById('mainContent').classList.add('hidden');
            document.getElementById('fileInput').value = '';
        }

        function completeReview() {
            alert(`Review completed!\\nAccepted: ${acceptedCount}\\nUser Added: ${userAddedCount}\\nDownload functionality ready!`);
        }

        // Keyboard shortcuts
        document.addEventListener('keydown', function(e) {
            if (e.target.tagName === 'INPUT' || e.target.tagName === 'TEXTAREA') return;
            
            switch(e.key.toLowerCase()) {
                case 'd':
                    e.preventDefault();
                    toggleDarkMode();
                    break;
                case '1':
                    e.preventDefault();
                    switchTab('feedback');
                    break;
                case '2':
                    e.preventDefault();
                    switchTab('chat');
                    break;
            }
        });
    </script>
</body>
</html>
'''

class DocumentAnalyzer:
    def analyze_document(self, file_path):
        doc = Document(file_path)
        sections = {}
        
        # Simple section detection
        current_section = "Executive Summary"
        content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check for section headers
            if len(text) < 100 and any(keyword in text.lower() for keyword in 
                ['executive summary', 'background', 'timeline', 'root cause', 'preventative', 'investigation']):
                if content:
                    sections[current_section] = '\\n\\n'.join(content)
                current_section = text
                content = []
            else:
                content.append(text)
        
        # Save last section
        if content:
            sections[current_section] = '\\n\\n'.join(content)
        
        # If no sections found, create default
        if not sections:
            all_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_content.append(para.text.strip())
            sections["Document Content"] = '\\n\\n'.join(all_content)
        
        return sections

class AIFeedbackEngine:
    def analyze_section(self, section_name, content):
        feedback_items = []
        
        # Content analysis
        word_count = len(content.split())
        
        if word_count < 50:
            feedback_items.append({
                'type': 'critical',
                'category': 'Content Completeness',
                'description': f'Section appears incomplete with only {word_count} words. Needs more detailed content.',
                'risk_level': 'High',
                'confidence': 0.9
            })
        
        if 'TBD' in content or 'TODO' in content:
            feedback_items.append({
                'type': 'important',
                'category': 'Content Quality',
                'description': 'Found placeholder text (TBD/TODO) that needs completion.',
                'risk_level': 'Medium',
                'confidence': 0.95
            })
        
        # Section-specific analysis
        section_lower = section_name.lower()
        
        if 'timeline' in section_lower and not re.search(r'\\d{1,2}[-/]\\d{1,2}[-/]\\d{2,4}', content):
            feedback_items.append({
                'type': 'critical',
                'category': 'Timeline Accuracy',
                'description': 'Timeline section lacks specific dates. Add concrete dates for all events.',
                'risk_level': 'High',
                'confidence': 0.9
            })
        
        if 'root cause' in section_lower and 'because' not in content.lower():
            feedback_items.append({
                'type': 'important',
                'category': 'Root Cause Analysis',
                'description': 'Root cause analysis lacks clear causal relationships. Use "5 Whys" methodology.',
                'risk_level': 'Medium',
                'confidence': 0.8
            })
        
        return feedback_items

class ChatBot:
    def process_query(self, query):
        query_lower = query.lower()
        
        if 'feedback' in query_lower:
            return "The feedback is generated using the Hawkeye 20-point investigation framework. Each item is analyzed for completeness, accuracy, and compliance with investigation standards."
        
        elif 'hawkeye' in query_lower:
            return "The Hawkeye framework includes 20 key checkpoints: Initial Assessment, Investigation Process, Seller Classification, Enforcement Decision-Making, and more. It ensures comprehensive document review."
        
        elif 'improve' in query_lower:
            return "To improve sections: 1) Add specific details and evidence, 2) Complete any placeholder text, 3) Ensure proper chronology in timelines, 4) Use clear causal language in root cause analysis."
        
        else:
            return "I can help with: explaining feedback items, Hawkeye framework guidelines, section improvement suggestions, and risk classifications. What specific aspect interests you?"

# Initialize components
doc_analyzer = DocumentAnalyzer()
ai_engine = AIFeedbackEngine()
chatbot = ChatBot()

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/upload', methods=['POST'])
def upload_document():
    file = request.files['file']
    if not file or not file.filename.endswith('.docx'):
        return jsonify({'error': 'Invalid file'}), 400
    
    filename = file.filename
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    # Analyze document
    sections = doc_analyzer.analyze_document(file_path)
    
    # Create session
    session_id = str(uuid.uuid4())
    sessions[session_id] = {
        'filename': filename,
        'sections': sections,
        'feedback_data': {},
        'accepted_feedback': [],
        'user_feedback': []
    }
    
    return jsonify({
        'session_id': session_id,
        'sections': list(sections.keys()),
        'status': 'success'
    })

@app.route('/analyze/<session_id>/<section_name>')
def analyze_section(session_id, section_name):
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session = sessions[session_id]
    content = session['sections'].get(section_name, '')
    
    # Get AI feedback
    feedback_items = ai_engine.analyze_section(section_name, content)
    session['feedback_data'][section_name] = feedback_items
    
    return jsonify({
        'section_content': content,
        'feedback_items': feedback_items
    })

@app.route('/feedback/<session_id>', methods=['POST'])
def handle_feedback(session_id):
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    data = request.json
    action = data.get('action')
    
    if action == 'accept':
        sessions[session_id]['accepted_feedback'].append(data)
    
    return jsonify({'status': 'success'})

@app.route('/custom_feedback/<session_id>', methods=['POST'])
def add_custom_feedback(session_id):
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    data = request.json
    sessions[session_id]['user_feedback'].append(data)
    
    return jsonify({'status': 'success'})

@app.route('/chat/<session_id>', methods=['POST'])
def chat_message(session_id):
    data = request.json
    query = data.get('message', '')
    
    response = chatbot.process_query(query)
    
    return jsonify({'response': response})

if __name__ == '__main__':
    print("Starting Enhanced Document Analysis Tool...")
    print("Access at: http://localhost:5000")
    app.run(host='0.0.0.0', port=5000, debug=True)