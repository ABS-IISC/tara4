import os
import json
import zipfile
import shutil
import uuid
from datetime import datetime
from docx import Document
from docx.shared import RGBColor, Pt
from lxml import etree

class DocumentProcessor:
    def __init__(self):
        self.temp_dirs = []

    def create_document_with_comments(self, original_path, comments_data, output_filename=None):
        """Create a Word document with proper comments"""
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"reviewed_document_{timestamp}.docx"

        try:
            # Use the advanced comment insertion method
            return self._create_with_xml_comments(original_path, comments_data, output_filename)
        except Exception as e:
            print(f"Advanced comment method failed: {e}")
            # Fallback to annotation method
            return self._create_with_annotations(original_path, comments_data, output_filename)

    def _create_with_xml_comments(self, original_path, comments_data, output_filename):
        """Create document with XML-based comments"""
        temp_dir = f"temp_{uuid.uuid4()}"
        self.temp_dirs.append(temp_dir)
        
        try:
            # Create a copy of the original document
            temp_docx = f"{temp_dir}_temp.docx"
            doc = Document(original_path)
            doc.save(temp_docx)
            
            # Extract the docx as a zip
            os.makedirs(temp_dir, exist_ok=True)
            with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Create comments.xml
            comments_xml = self._generate_comments_xml(comments_data)
            comments_path = os.path.join(temp_dir, 'word', 'comments.xml')
            
            with open(comments_path, 'w', encoding='utf-8') as f:
                f.write(comments_xml)
            
            # Update document.xml.rels
            self._update_document_rels(temp_dir)
            
            # Update [Content_Types].xml
            self._update_content_types(temp_dir)
            
            # Update document.xml with comment references
            self._insert_comment_references(temp_dir, comments_data)
            
            # Repackage as docx
            with zipfile.ZipFile(output_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, temp_dir)
                        zipf.write(file_path, arcname)
            
            return output_filename
            
        finally:
            # Cleanup
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

    def _generate_comments_xml(self, comments_data):
        """Generate the comments.xml content"""
        xml_content = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'''
        
        for i, comment in enumerate(comments_data):
            comment_id = i + 1
            author = comment.get('author', 'AI Feedback')
            date = datetime.now().strftime('%Y-%m-%dT%H:%M:%S.%fZ')
            text = comment.get('comment', '')
            
            xml_content += f'''
    <w:comment w:id="{comment_id}" w:author="{author}" w:date="{date}">
        <w:p>
            <w:r>
                <w:t>{text}</w:t>
            </w:r>
        </w:p>
    </w:comment>'''
        
        xml_content += '\n</w:comments>'
        return xml_content

    def _update_document_rels(self, temp_dir):
        """Update document.xml.rels to include comments relationship"""
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        
        if os.path.exists(rels_path):
            with open(rels_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if 'comments.xml' not in content:
                # Find the highest rId number
                import re
                rids = re.findall(r'rId(\d+)', content)
                max_rid = max([int(rid) for rid in rids]) if rids else 0
                new_rid = f"rId{max_rid + 1}"
                
                new_rel = f'<Relationship Id="{new_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                content = content.replace('</Relationships>', f'{new_rel}</Relationships>')
                
                with open(rels_path, 'w', encoding='utf-8') as f:
                    f.write(content)

    def _update_content_types(self, temp_dir):
        """Update [Content_Types].xml to include comments content type"""
        content_types_path = os.path.join(temp_dir, '[Content_Types].xml')
        
        if os.path.exists(content_types_path):
            with open(content_types_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if 'comments.xml' not in content:
                new_type = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                content = content.replace('</Types>', f'{new_type}</Types>')
                
                with open(content_types_path, 'w', encoding='utf-8') as f:
                    f.write(content)

    def _insert_comment_references(self, temp_dir, comments_data):
        """Insert comment references into document.xml"""
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        
        if not os.path.exists(doc_xml_path):
            return
        
        # Parse the document XML
        tree = etree.parse(doc_xml_path)
        root = tree.getroot()
        
        # Find paragraphs and insert comment references
        paragraphs = root.xpath('//w:p', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        for i, comment in enumerate(comments_data):
            comment_id = i + 1
            para_index = comment.get('paragraph_index', 0)
            
            if para_index < len(paragraphs):
                para = paragraphs[para_index]
                
                # Create comment range start
                comment_start = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
                comment_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                
                # Create comment range end
                comment_end = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
                comment_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                
                # Create comment reference
                comment_ref = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                comment_ref_elem = etree.SubElement(comment_ref, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference')
                comment_ref_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                
                # Insert at the beginning of the paragraph
                para.insert(0, comment_start)
                para.append(comment_end)
                para.append(comment_ref)
        
        # Save the modified document.xml
        tree.write(doc_xml_path, encoding='utf-8', xml_declaration=True)

    def _create_with_annotations(self, original_path, comments_data, output_filename):
        """Fallback method: create document with inline annotations"""
        try:
            original_doc = Document(original_path)
            new_doc = Document()
            
            # Copy document structure and add annotations
            for para_idx, para in enumerate(original_doc.paragraphs):
                # Copy paragraph
                new_para = new_doc.add_paragraph()
                new_para.style = para.style
                
                # Copy runs
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    if run.bold:
                        new_run.bold = True
                    if run.italic:
                        new_run.italic = True
                    if run.underline:
                        new_run.underline = True
                
                # Add comments for this paragraph
                para_comments = [c for c in comments_data if c.get('paragraph_index') == para_idx]
                
                for comment in para_comments:
                    comment_para = new_doc.add_paragraph()
                    comment_run = comment_para.add_run(f"💬 {comment.get('author', 'AI Feedback')}: {comment.get('comment', '')}")
                    comment_run.font.color.rgb = RGBColor(102, 126, 234)
                    comment_run.font.size = Pt(10)
                    comment_run.italic = True
            
            # Add summary section
            self._add_feedback_summary(new_doc, comments_data)
            
            new_doc.save(output_filename)
            return output_filename
            
        except Exception as e:
            print(f"Annotation method failed: {e}")
            return self._create_simple_copy(original_path, comments_data, output_filename)

    def _create_simple_copy(self, original_path, comments_data, output_filename):
        """Simple fallback: copy document and add summary"""
        try:
            doc = Document(original_path)
            
            # Add feedback summary at the end
            self._add_feedback_summary(doc, comments_data)
            
            doc.save(output_filename)
            return output_filename
            
        except Exception as e:
            print(f"Simple copy method failed: {e}")
            return None

    def _add_feedback_summary(self, doc, comments_data):
        """Add comprehensive feedback summary to document"""
        doc.add_page_break()
        
        # Title
        title = doc.add_heading('Hawkeye Review Feedback Summary', 1)
        
        # Metadata
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Total feedback items: {len(comments_data)}')
        doc.add_paragraph('')
        
        # Group comments by section
        from collections import defaultdict
        section_comments = defaultdict(list)
        
        for comment in comments_data:
            section = comment.get('section', 'Unknown Section')
            section_comments[section].append(comment)
        
        # Add each section's feedback
        for section_name, comments in section_comments.items():
            section_heading = doc.add_heading(section_name, 2)
            
            for i, comment in enumerate(comments, 1):
                # Create feedback item
                para = doc.add_paragraph(style='List Number')
                
                # Author and type
                author_run = para.add_run(f"[{comment.get('author', 'AI Feedback')}] ")
                author_run.bold = True
                
                # Risk level with color
                risk_level = comment.get('risk_level', 'Low')
                type_text = f"{comment.get('type', 'feedback').upper()} - {risk_level} Risk: "
                type_run = para.add_run(type_text)
                type_run.bold = True
                
                if risk_level == 'High':
                    type_run.font.color.rgb = RGBColor(231, 76, 60)  # Red
                elif risk_level == 'Medium':
                    type_run.font.color.rgb = RGBColor(243, 156, 18)  # Orange
                else:
                    type_run.font.color.rgb = RGBColor(52, 152, 219)  # Blue
                
                # Comment text
                para.add_run(comment.get('comment', ''))
                
                # Add spacing
                if i < len(comments):
                    doc.add_paragraph('')

    def cleanup_temp_files(self):
        """Clean up temporary directories"""
        for temp_dir in self.temp_dirs:
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                except:
                    pass
        self.temp_dirs = []

    def __del__(self):
        """Cleanup on destruction"""
        self.cleanup_temp_files()